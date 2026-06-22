import io
import logging
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document
from src.services.database import db
from src.services.gemini_client import generate_embedding
from src.config import CHUNK_SIZE, CHUNK_OVERLAP

logger = logging.getLogger(__name__)

class IngestionAgent:
    def __init__(self):
        pass

    def extract_text(self, file_bytes: bytes, file_name: str) -> str:
        """Extracts plain text from pdf or docx bytes."""
        extension = file_name.split('.')[-1].lower()
        text = ""
        
        try:
            if extension == 'pdf':
                logger.info(f"Extracting PDF text from {file_name}...")
                pdf_file = io.BytesIO(file_bytes)
                reader = PdfReader(pdf_file)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
                
            elif extension == 'docx':
                logger.info(f"Extracting DOCX text from {file_name}...")
                docx_file = io.BytesIO(file_bytes)
                doc = Document(docx_file)
                paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also read tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            tables_text.append(row_text)
                text = "\n".join(paragraphs_text + tables_text)
                
            else:
                logger.warning(f"Unsupported file format: {extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {e}")
            raise ValueError(f"Failed to parse file content: {str(e)}")
            
        return text

    def chunk_text(self, text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
        """Splits text into chunks of chunk_size characters with overlap."""
        chunks = []
        text_len = len(text)
        
        if text_len <= chunk_size:
            if text.strip():
                return [text.strip()]
            return []
            
        start = 0
        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            # Advance start pointer by difference (chunk_size - overlap)
            start += (chunk_size - overlap)
            
        return chunks

    def ingest_file(self, session_id: str, file_id: str, file_name: str, file_size: str, file_type: str, file_bytes: bytes) -> Dict[str, Any]:
        """Orchestrates the ingestion workflow for a single document."""
        # 1. Extract text
        raw_text = self.extract_text(file_bytes, file_name)
        if not raw_text.strip():
            logger.warning(f"No text extracted from {file_name}.")
            raw_text = "[Empty Document Content]"

        # 2. Add document metadata to DB with its size in bytes
        db.add_document(session_id, file_id, file_name, file_size, file_type, raw_text, len(file_bytes), file_bytes)

        # 3. Create chunks
        chunks = self.chunk_text(raw_text)
        logger.info(f"Split {file_name} into {len(chunks)} chunks.")

        # 4. Generate embeddings and save chunks
        for i, chunk in enumerate(chunks):
            # We can log progress periodically
            if i % 10 == 0:
                logger.info(f"Generating embedding for chunk {i}/{len(chunks)} of {file_name}...")
            
            embedding = generate_embedding(chunk)
            db.add_chunk(session_id, file_id, chunk, embedding)

        logger.info(f"Successfully ingested document {file_name} (ID: {file_id}) for session: {session_id}")
        return {
            "id": file_id,
            "name": file_name,
            "size": file_size,
            "type": file_type,
            "chunks_count": len(chunks)
        }

ingestion_agent = IngestionAgent()
