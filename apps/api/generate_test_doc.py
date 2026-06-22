import os
from docx import Document

def generate():
    doc = Document()
    doc.add_heading('AIRA Project Test Document', level=0)
    
    doc.add_heading('1. Market Landscape & Capital Activity', level=1)
    doc.add_paragraph(
        'According to the Market_Analysis_2026 report, there is a strong shift towards agentic system platforms in Q1 2026. '
        'Capital allocation has increased by 45% compared to Q4 2025.'
    )
    
    doc.add_heading('2. Operational Milestones', level=1)
    doc.add_paragraph(
        'The project milestones track a consistent delivery velocity. The core LLM pipelines were completed 2 weeks ahead of schedule.'
    )
    
    doc.add_heading('3. Identified Bottlenecks', level=1)
    doc.add_paragraph(
        'Deployment latency in secondary staging environments is currently lagging by 12% due to integration overhead.'
    )
    
    dest_path = 'c:/Users/wesam/AIRA/apps/api/test_doc.docx'
    doc.save(dest_path)
    print(f"Test document saved to {dest_path}")

if __name__ == '__main__':
    generate()
